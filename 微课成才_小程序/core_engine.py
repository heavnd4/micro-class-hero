import json
import subprocess
import os
import time
import gc
import re
from pathlib import Path
import anthropic
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import cv2
import numpy as np
from paddleocr import PaddleOCR

class VideoProcessor:
    def __init__(self, base_dir=None, api_key=None):
        # 1. 自动计算路径 (兼容本地 E 盘和云端环境)
        if base_dir:
            self.base_dir = Path(base_dir)
        else:
            # 默认假设在小程序目录下，向上两级到项目根目录
            self.base_dir = Path(__file__).parent.parent
            
        self.models_dir = self.base_dir / "models"
        
        # 2. 初始化配置
        self.api_key = api_key or "sk-sp-a057b4bf1def40ad8b44d6a908b59a84"
        self.base_url = "https://coding.dashscope.aliyuncs.com/apps/anthropic"
        self.model_name = "qwen3.5-plus"
        
        # 3. 环境变量强控 (针对 PaddleOCR)
        os.environ["PADDLE_HOME"] = str(self.models_dir / ".paddleocr")
        os.environ["PADDLE_PDX_HOME"] = str(self.models_dir / ".paddlex")
        os.environ["FLAGS_use_mkldnn"] = "0"
        os.environ["FLAGS_enable_pir_api"] = "0"
        
        # 4. 初始化 AI 客户端
        self.client = anthropic.Anthropic(
            api_key=self.api_key,
            base_url=self.base_url
        )
        
        self.stop_requested = False
        self.progress = 0
        self.current_step = "等待中"

    def get_status_manager(self, out_dir: Path):
        status_file = out_dir / "task_status.json"
        def load():
            if status_file.exists():
                try:
                    return json.loads(status_file.read_text(encoding="utf-8"))
                except:
                    pass
            return {}
        def save(status: dict):
            status_file.write_text(json.dumps(status, ensure_ascii=False, indent=2), encoding="utf-8")
        return load, save

    def transcribe(self, video_path: Path, out_dir: Path):
        """第一步：语音转文字"""
        self.current_step = "语音转录中"
        self.progress = 0
        txt_file = out_dir / "01_transcript.txt"
        if txt_file.exists(): 
            print("  ✓ 转录文稿已存在，跳过")
            return txt_file.read_text(encoding="utf-8")

        from faster_whisper import WhisperModel
        whisper_path = self.models_dir / "models--Systran--faster-whisper-small" / "snapshots" / "main"
        # 显式使用 CPU 和 INT8 模式，确保稳定性
        model = WhisperModel(str(whisper_path), device="cpu", compute_type="int8")
        segments, _ = model.transcribe(str(video_path), language="zh")
        
        transcript_lines = []
        for seg in segments:
            transcript_lines.append(seg.text.strip())
        
        full_text = "\n".join(transcript_lines)
        txt_file.write_text(full_text, encoding="utf-8")
        return full_text

    def extract_frames(self, video_path: Path, out_dir: Path, mode="smart"):
        """第二步：关键帧提取"""
        self.current_step = "截取关键帧"
        frames_dir = out_dir / "frames"
        frames_dir.mkdir(exist_ok=True)
        existing = sorted(frames_dir.glob("frame_*.jpg"))
        if existing: 
            print(f"  ✓ 截图已存在 ({len(existing)} 张)，跳过")
            return existing

        if mode == "fixed":
            subprocess.run(["ffmpeg", "-i", str(video_path), "-vf", "fps=1/60", "-q:v", "2", str(frames_dir / "frame_%03d.jpg"), "-y"], capture_output=True)
        else:
            temp_pattern = str(frames_dir / "raw_%04d.jpg")
            subprocess.run(["ffmpeg", "-i", str(video_path), "-vf", "select='gt(scene,0.005)',setpts=N/FRAME_RATE/TB", "-vsync", "vfr", "-q:v", "2", temp_pattern, "-y"], capture_output=True)
            
            raw_frames = sorted(frames_dir.glob("raw_*.jpg"))
            if not raw_frames: return []
            
            ocr = PaddleOCR(use_angle_cls=True, lang='ch')
            final_frames = []
            last_text_set = set()
            
            total = len(raw_frames)
            for i, frame_path in enumerate(raw_frames, 1):
                if self.stop_requested: raise InterruptedError("Stopped by user")
                self.progress = int((i / total) * 100)
                
                result = ocr.ocr(str(frame_path))
                current_text = "".join([line[1][0] for line in result[0]]) if result and result[0] else ""
                current_text_set = set(current_text)
                
                keep = (i == 1 and current_text) or (current_text and (not last_text_set or len(current_text_set.intersection(last_text_set))/max(len(current_text_set), 1) < 0.8))
                
                if keep:
                    new_path = frames_dir / f"frame_{len(final_frames)+1:03d}.jpg"
                    os.rename(str(frame_path), str(new_path))
                    final_frames.append(new_path)
                    last_text_set = current_text_set
                else:
                    os.remove(str(frame_path))
                gc.collect()
        return sorted(frames_dir.glob("frame_*.jpg"))

    def structure_text(self, transcript: str, out_dir: Path):
        """第三步：AI 整理文稿结构"""
        self.current_step = "AI 结构化整理"
        self.progress = 50
        cache_file = out_dir / "02_structured.json"
        if cache_file.exists(): return json.loads(cache_file.read_text(encoding="utf-8"))

        text_input = transcript[:8000]
        prompt = f"你是一位专业的课程内容整理专家。请将以下原始字幕划分为5-8个章节，改写为书面化正文并以JSON格式返回。JSON格式包含'title'和'chapters'(index, title, content)：\n\n{text_input}"
        
        response = self.client.messages.create(
            model=self.model_name,
            max_tokens=4000,
            messages=[{"role": "user", "content": prompt}]
        )
        
        raw = "".join([block.text for block in response.content if hasattr(block, "text")]).strip()
        if "```json" in raw: raw = raw.split("```json")[1].split("```")[0].strip()
        elif "```" in raw: raw = raw.split("```")[1].split("```")[0].strip()
        
        data = json.loads(raw)
        cache_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        return data

    def generate_questions(self, structured: dict, out_dir: Path):
        """第四步：AI 生成考题"""
        self.current_step = "AI 生成考题"
        self.progress = 80
        cache_file = out_dir / "03_questions.json"
        if cache_file.exists(): return json.loads(cache_file.read_text(encoding="utf-8"))

        full_content = "\n\n".join([f"{ch['title']}\n{ch['content']}" for ch in structured["chapters"]])
        prompt = f"你是一位专业的教育测评专家。请根据以下内容编写一套配套测试题（25道左右，含单选、多选、判断、简答），以JSON数组格式返回：\n\n{full_content}"
        
        response = self.client.messages.create(
            model=self.model_name,
            max_tokens=8000,
            messages=[{"role": "user", "content": prompt}]
        )
        
        raw = "".join([block.text for block in response.content if hasattr(block, "text")]).strip()
        match = re.search(r'\[.*\]', raw, re.DOTALL)
        if match:
            all_questions = json.loads(match.group())
            cache_file.write_text(json.dumps(all_questions, ensure_ascii=False, indent=2), encoding="utf-8")
            return all_questions
        return []

    def save_docx(self, structured, frames, questions, out_dir, video_stem):
        """第五步：生成 Word 文档"""
        self.current_step = "生成 Word 文档"
        self.progress = 95
        # 1. 教材
        doc = Document()
        doc.add_heading(structured.get("title", video_stem), 0)
        for ch in structured["chapters"]:
            doc.add_heading(f"第{ch['index']}章 {ch['title']}", level=1)
            doc.add_paragraph(ch["content"])
            if frames:
                idx = int((ch["index"]-1) * len(frames) / len(structured["chapters"]))
                if idx < len(frames):
                    doc.add_picture(str(frames[idx]), width=Inches(5))
        
        timestamp = time.strftime("%H%M%S")
        book_path = out_dir / f"{video_stem}_教材_{timestamp}.docx"
        doc.save(str(book_path))

        # 2. 考题
        q_doc = Document()
        q_doc.add_heading(f"《{video_stem}》测试题库", 0)
        for i, q in enumerate(questions, 1):
            q_doc.add_paragraph(f"{i}. {q['question']} ({q['type']})")
            if "options" in q:
                for opt in q["options"]: q_doc.add_paragraph(f"   {opt}")
        
        q_path = out_dir / f"{video_stem}_考题_{timestamp}.docx"
        q_doc.save(str(q_path))
        return book_path, q_path

    def process(self, video_path: str, out_root: str):
        video_path = Path(video_path)
        out_dir = Path(out_root) / video_path.stem
        out_dir.mkdir(parents=True, exist_ok=True)
        
        load_status, save_status = self.get_status_manager(out_dir)
        status = load_status()

        print(f"开始处理: {video_path.stem}")
        
        # 1. 转录
        transcript = self.transcribe(video_path, out_dir)
        # 2. 截图
        frames = self.extract_frames(video_path, out_dir)
        # 3. 结构化
        structured = self.structure_text(transcript, out_dir)
        # 4. 考题
        questions = self.generate_questions(structured, out_dir)
        # 5. 保存
        self.save_docx(structured, frames, questions, out_dir, video_path.stem)
        
        self.current_step = "已完成"
        self.progress = 100
        print(f"🎉 处理圆满成功!")

if __name__ == "__main__":
    p = VideoProcessor(base_dir=r"E:\.cc项目")
    p.process(r"E:\.cc项目\微课成才_小程序\temp_input\test_video.mp4", r"E:\.cc项目\微课成才_小程序\temp_output")
