"""
视频 → Word 产品生成器（单视频版）
输出：教材.docx + 考题.docx
"""

import json
import subprocess
from pathlib import Path
import anthropic
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================
# ★ 配置区
# ============================================================
# 待处理视频所在的文件夹
VIDEO_DIR     = r"E:\.cc项目\视频转换\输入"
# 结果输出的总文件夹
OUTPUT_ROOT   = r"E:\.cc项目\视频转换\输出"

API_KEY       = "sk-sp-a057b4bf1def40ad8b44d6a908b59a84"
BASE_URL      = "https://coding.dashscope.aliyuncs.com/apps/anthropic"
MODEL         = "qwen3.5-plus"
FRAME_EVERY   = 60
WHISPER_MODEL = "small"
# 截图模式：'smart' (智能识别) 或 'fixed' (fixed)
EXTRACT_MODE  = "smart"
# 停止请求标志 (2026-04-24 新增)
STOP_REQUESTED = False
# ============================================================


import os
import time
import cv2
import numpy as np
from paddleocr import PaddleOCR

# 设置 PaddleOCR 模型存储路径，避免占用 C 盘
os.environ["PADDLE_HOME"] = r"E:\.cc项目\models\.paddleocr"
os.environ["PADDLE_PDX_HOME"] = r"E:\.cc项目\models\.paddlex"
# 禁用 MKLDNN 以解决 (Unimplemented) ConvertPirAttribute2RuntimeAttribute 报错
os.environ["FLAGS_use_mkldnn"] = "0"
os.environ["FLAGS_enable_pir_api"] = "0"

def get_status_manager(out_dir: Path):
    """状态管理器：记录步骤完成情况"""
    status_file = out_dir / "task_status.json"
    
    def load():
        if status_file.exists():
            try:
                return json.loads(status_file.read_text(encoding="utf-8"))
            except:
                pass
        return {}

    def save(status: dict):
        status_file.write_text(json.dumps(status, ensure_ascii=False, indent=2), encoding="utf-8")

    return load, save


def format_timestamp(seconds: float) -> str:
    """将秒数转换为 SRT 时间戳格式 (00:00:00,000)"""
    td = time.gmtime(seconds)
    milli = int((seconds - int(seconds)) * 1000)
    return f"{time.strftime('%H:%M:%S', td)},{milli:03d}"


def step1_transcribe(video_path: Path, out_dir: Path) -> str:
    """第一步：视频转文字 (同步生成 txt 和 srt)"""
    txt_file = out_dir / "01_transcript.txt"
    srt_file = out_dir / f"{video_path.stem}_字幕.srt"

    if txt_file.exists() and srt_file.exists():
        print("  ✓ 转录文稿及字幕已存在，跳过")
        return txt_file.read_text(encoding="utf-8")

    print("  → 开始转录，请耐心等待（30分钟视频约需15-25分钟）...")
    print("     转录期间可以做其他事，完成后会提示")

    from faster_whisper import WhisperModel
    model = WhisperModel(r"E:\.cc项目\models\models--Systran--faster-whisper-small\snapshots\main", device="cpu", compute_type="int8")
    segments, _ = model.transcribe(str(video_path), language="zh")

    transcript_lines = []
    srt_lines = []
    
    for i, seg in enumerate(segments, 1):
        text = seg.text.strip()
        transcript_lines.append(text)
        
        # 构建 SRT 格式
        start_time = format_timestamp(seg.start)
        end_time = format_timestamp(seg.end)
        srt_lines.append(f"{i}\n{start_time} --> {end_time}\n{text}\n")

    full_text = "\n".join(transcript_lines)
    txt_file.write_text(full_text, encoding="utf-8")
    srt_file.write_text("\n".join(srt_lines), encoding="utf-8")
    
    print(f"  ✓ 转录完成，共 {len(full_text)} 字，字幕文件已生成")
    return full_text


def step2_extract_frames(video_path: Path, out_dir: Path) -> list:
    """第二步：智能/固定截图切换"""
    frames_dir = out_dir / "frames"
    frames_dir.mkdir(exist_ok=True)

    existing = sorted(frames_dir.glob("frame_*.jpg"))
    if existing:
        print(f"  ✓ 截图已存在（{len(existing)} 张），跳过")
        return existing

    # --- 模式判断：固定频率模式 ---
    if EXTRACT_MODE == "fixed":
        print(f"  → 截图中（固定频率：每 {FRAME_EVERY} 秒截一帧）...")
        output_pattern = str(frames_dir / "frame_%03d.jpg")
        cmd = [
            "ffmpeg", "-i", str(video_path),
            "-vf", f"fps=1/{FRAME_EVERY}",
            "-q:v", "2",
            output_pattern, "-y"
        ]
        subprocess.run(cmd, capture_output=True)
        return sorted(frames_dir.glob("frame_*.jpg"))

    # --- 模式判断：智能识别模式 ---
    print("  → 智能截图中（识别 PPT 翻页）...")
    
    # 1. 使用 FFmpeg 的场景检测初步提取变动帧
    temp_pattern = str(frames_dir / "raw_%04d.jpg")
    # gt(scene,0.005) 比较灵敏，捕捉大多数翻页
    cmd = [
        "ffmpeg", "-i", str(video_path),
        "-vf", "select='gt(scene,0.005)',setpts=N/FRAME_RATE/TB",
        "-vsync", "vfr", "-q:v", "2",
        temp_pattern, "-y"
    ]
    subprocess.run(cmd, capture_output=True, text=True, errors='ignore')

    raw_frames = sorted(frames_dir.glob("raw_*.jpg"))
    if not raw_frames:
        print("    ! 场景检测未捕获到帧，回退到固定频率模式")
        output_pattern = str(frames_dir / "frame_%03d.jpg")
        cmd = ["ffmpeg", "-i", str(video_path), "-vf", f"fps=1/{FRAME_EVERY}", "-q:v", "2", output_pattern, "-y"]
        subprocess.run(cmd, capture_output=True)
        return sorted(frames_dir.glob("frame_*.jpg"))

    # 2. OCR 二次过滤，剔除重复文本帧
    total_raw = len(raw_frames)
    print(f"    初选 {total_raw} 张，正在进行 OCR 语义去重...")
    # 只在需要时初始化 OCR，显式禁用所有加速以防报错
    ocr = PaddleOCR(
        use_textline_orientation=True, 
        lang='ch'
    )
    
    final_frames = []
    last_text_set = set()
    import gc

    for i, frame_path in enumerate(raw_frames, 1):
        # --- 中断检查点 ---
        if STOP_REQUESTED:
            raise InterruptedError("用户手动终止了任务")

        if i % 10 == 0 or i == 1:
            print(f"      → 进度: [{i}/{total_raw}]...")
            
        try:
            # 识别文字
            result = ocr.ocr(str(frame_path))
            
            current_text = ""
            if result and result[0]:
                texts = [line[1][0] for line in result[0] if line[1][1] > 0.6]
                current_text = "".join(texts)
            
            current_text_set = set(current_text)
            
            keep = False
            if i == 1:
                keep = len(current_text_set) > 0 
            else:
                if len(current_text_set) == 0:
                    keep = False
                elif len(last_text_set) == 0:
                    keep = True
                else:
                    intersection = current_text_set.intersection(last_text_set)
                    similarity = len(intersection) / max(len(current_text_set), len(last_text_set))
                    keep = similarity < 0.8

            if keep:
                new_path = frames_dir / f"frame_{len(final_frames)+1:03d}.jpg"
                if new_path.exists(): os.remove(str(new_path))
                os.rename(str(frame_path), str(new_path))
                final_frames.append(new_path)
                last_text_set = current_text_set
            else:
                if os.path.exists(str(frame_path)):
                    os.remove(str(frame_path))
            
            # 每帧清理一次内存
            gc.collect()

        except Exception as e:
            print(f"    ! 处理第 {i} 帧时出错: {e}")
            if os.path.exists(str(frame_path)): os.remove(str(frame_path))

    # 清理所有剩余的 raw_ 帧
    for f in frames_dir.glob("raw_*.jpg"):
        try: os.remove(str(f))
        except: pass

    if not final_frames:
         print("    ! 智能去重后无剩余帧，回退到固定频率模式")
         output_pattern = str(frames_dir / "frame_%03d.jpg")
         cmd = ["ffmpeg", "-i", str(video_path), "-vf", f"fps=1/{FRAME_EVERY}", "-q:v", "2", output_pattern, "-y"]
         subprocess.run(cmd, capture_output=True)
         final_frames = sorted(frames_dir.glob("frame_*.jpg"))

    print(f"  ✓ 智能截图完成，精选 {len(final_frames)} 张关键帧")
    return final_frames


def step3_structure_text(transcript: str, out_dir: Path, client) -> dict:
    """第三步：AI 整理文稿结构"""
    cache_file = out_dir / "02_structured.json"

    if cache_file.exists():
        print("  ✓ 结构化数据已存在，跳过")
        return json.loads(cache_file.read_text(encoding="utf-8"))

    print("  → 调用 AI 整理文稿，约需1-2分钟...")

    text_input = transcript[:8000] if len(transcript) > 8000 else transcript

    prompt = f"""你是一位专业的课程内容整理专家。
以下是一段技术视频的原始字幕，请完成以下任务：

1. 将内容划分为5-8个章节，每章节有简洁的标题
2. 每章节内容改写为书面化正文（去除"嗯""那个"等口语，保留所有技术术语和操作步骤）
3. 每章节正文200-400字

请用以下JSON格式返回（只返回JSON，不要其他内容）：
{{
  "title": "根据内容总结的课程标题",
  "chapters": [
    {{
      "index": 1,
      "title": "章节标题",
      "content": "书面化正文内容..."
    }}
  ]
}}

原始字幕：
{text_input}"""

    response = client.messages.create(
        model=MODEL,
        max_tokens=4000,
        messages=[{"role": "user", "content": prompt}]
    )

    # 提取文本内容，兼容多种 Block 类型（如 TextBlock, ThinkingBlock）
    raw = ""
    for block in response.content:
        if hasattr(block, "text"):
            raw += block.text
    raw = raw.strip()
    if "```" in raw:
        parts = raw.split("```")
        for part in parts:
            part = part.strip()
            if part.startswith("json"):
                part = part[4:].strip()
            if part.startswith("{"):
                raw = part
                break

    data = json.loads(raw)
    cache_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"  ✓ 整理完成，共 {len(data['chapters'])} 章")
    return data


def step4_generate_questions(structured: dict, out_dir: Path, video_stem: str, client) -> list:
    """第四步：AI 生成考题 (一次性生成全库以保证数量和质量)"""
    cache_file = out_dir / "03_questions.json"

    if cache_file.exists():
        print("  ✓ 考题数据已存在，跳过")
        return json.loads(cache_file.read_text(encoding="utf-8"))

    print("  → 正在根据全文生成考题库，约需1-2分钟...")
    
    # 汇总全文内容
    full_content = "\n\n".join([f"章节{ch['index']}: {ch['title']}\n{ch['content']}" for ch in structured["chapters"]])
    
    prompt = f"""你是一位专业的教育测评专家。请根据以下课程内容，编写一套配套测试题。

课程标题：{structured.get('title', video_stem)}

要求：
1. 总题量：25-27道。
2. 题型分布：
   - 单选题：15道
   - 多选题：5道
   - 判断题：5道
   - 简答题：1-2道
3. 内容覆盖：题目必须覆盖课程的所有主要技术点（电路、插座、给排水、燃气等）。
4. 难度适中，解析详尽。

只返回一个JSON数组，格式如下：
[
  {{"type": "单选题", "question": "...", "options": ["A.", "B.", "C.", "D."], "answer": "A", "explanation": "..."}},
  {{"type": "多选题", "question": "...", "options": ["A.", "B.", "C.", "D."], "answer": "AB", "explanation": "..."}},
  {{"type": "判断题", "question": "...", "answer": "正确", "explanation": "..."}},
  {{"type": "简答题", "question": "...", "answer": "..."}}
]

课程内容：
{full_content}"""

    response = client.messages.create(
        model=MODEL,
        max_tokens=8000,
        messages=[{"role": "user", "content": prompt}]
    )

    # 提取文本内容
    raw = ""
    for block in response.content:
        if hasattr(block, "text"):
            raw += block.text
    raw = raw.strip()
    
    # 提取 JSON 数组部分
    import re
    match = re.search(r'\[.*\]', raw, re.DOTALL)
    if match:
        json_str = match.group()
        try:
            all_questions = json.loads(json_str)
        except json.JSONDecodeError:
            print("  ✗ JSON 解析失败，尝试修复并重新解析...")
            # 简单的修复逻辑：如果是因为截断导致的
            if not json_str.endswith("]"):
                json_str += "]"
            try:
                all_questions = json.loads(json_str)
            except:
                print("  ✗ 无法自动修复 JSON，请检查 API 输出")
                all_questions = []
    else:
        print("  ✗ 未在 AI 响应中找到 JSON 数组")
        all_questions = []

    cache_file.write_text(json.dumps(all_questions, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"  ✓ 考题生成完成，共 {len(all_questions)} 道题")
    return all_questions


def make_book_docx(structured: dict, frames: list, out_path: Path, video_stem: str):
    """生成图文教材 Word"""
    doc = Document()

    section = doc.sections[0]
    section.page_width    = Inches(8.27)
    section.page_height   = Inches(11.69)
    section.left_margin   = Inches(1.18)
    section.right_margin  = Inches(1.18)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # 封面
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    run = p.add_run(structured.get("title", video_stem))
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2C, 0x4A, 0x8C)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"\n课程：{video_stem}")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()

    # 目录
    p_toc = doc.add_paragraph()
    run_toc = p_toc.add_run("目  录")
    run_toc.font.size = Pt(18)
    run_toc.font.bold = True
    doc.add_paragraph()
    for ch in structured["chapters"]:
        p = doc.add_paragraph(f"第{ch['index']}章  {ch['title']}")
        p.paragraph_format.left_indent = Pt(20)
    doc.add_page_break()

    # 正文
    n_chapters = len(structured["chapters"])
    n_frames   = len(frames)

    for ch in structured["chapters"]:
        p_title = doc.add_paragraph()
        run_title = p_title.add_run(f"第{ch['index']}章  {ch['title']}")
        run_title.font.size = Pt(16)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(0x2C, 0x4A, 0x8C)
        p_title.paragraph_format.space_before = Pt(12)
        p_title.paragraph_format.space_after  = Pt(6)

        p_body = doc.add_paragraph(ch["content"])
        p_body.paragraph_format.first_line_indent = Pt(28)

        if n_frames > 0:
            start_idx = int((ch["index"] - 1) / n_chapters * n_frames)
            end_idx   = int(ch["index"] / n_chapters * n_frames)
            chapter_frames = frames[start_idx:end_idx]

            for i, img_path in enumerate(chapter_frames[:2]):
                try:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_img = p_img.add_run()
                    run_img.add_picture(str(img_path), width=Inches(5.2))

                    p_cap = doc.add_paragraph(f"图 {ch['index']}-{i+1}  {ch['title']}")
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.runs[0].font.size = Pt(9)
                    p_cap.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                except Exception as e:
                    print(f"     插图失败（{img_path.name}）：{e}")

        doc.add_paragraph()

    doc.save(str(out_path))
    print(f"  ✓ 教材已保存 → {out_path.name}")


def make_questions_docx(structured: dict, all_questions: list, out_path: Path, video_stem: str):
    """生成考题 Word"""
    doc = Document()

    section = doc.sections[0]
    section.page_width    = Inches(8.27)
    section.page_height   = Inches(11.69)
    section.left_margin   = Inches(1.18)
    section.right_margin  = Inches(1.18)

    # 封面
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run(f"《{structured.get('title', video_stem)}》\n配套测试题库")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2C, 0x4A, 0x8C)
    doc.add_page_break()

    q_num = 1
    all_answers = []

    # 按题型分类列出所有题目
    for qtype in ["单选题", "多选题", "判断题", "简答题"]:
        qs = [q for q in all_questions if q.get("type") == qtype]
        if not qs:
            continue

        p_type = doc.add_paragraph(f"【{qtype}】")
        p_type.runs[0].font.size = Pt(14)
        p_type.runs[0].font.bold = True
        p_type.runs[0].font.color.rgb = RGBColor(0x1A, 0x6B, 0x4A)
        doc.add_paragraph()

        for q in qs:
            p_q = doc.add_paragraph(f"{q_num}. {q['question']}")
            p_q.runs[0].font.size = Pt(11)

            if "options" in q:
                for opt in q["options"]:
                    p_opt = doc.add_paragraph(opt)
                    p_opt.paragraph_format.left_indent = Pt(28)

            all_answers.append({
                "num": q_num,
                "answer": q.get("answer", ""),
                "explanation": q.get("explanation", "")
            })
            q_num += 1
        doc.add_paragraph()

    # 答案页
    doc.add_page_break()
    p_ans_head = doc.add_paragraph("▶ 参考答案及解析")
    p_ans_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ans_head.runs[0].font.size = Pt(16)
    p_ans_head.runs[0].font.bold = True
    p_ans_head.runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    doc.add_paragraph()

    for ans in all_answers:
        ans_text = f"{ans['num']}. {ans['answer']}"
        if ans["explanation"]:
            ans_text += f"　【解析】{ans['explanation']}"
        p_a = doc.add_paragraph(ans_text)
        p_a.paragraph_format.left_indent = Pt(14)
        p_a.runs[0].font.size = Pt(10)
        p_a.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.save(str(out_path))
    print(f"  ✓ 考题已保存 → {out_path.name}")


def process_single_video(video_path: Path, out_root: Path, client):
    """处理单个视频的完整流程"""
    video_stem = video_path.stem
    out_dir = out_root / video_stem
    out_dir.mkdir(parents=True, exist_ok=True)
    
    print(f"\n{'='*10} 正在处理: {video_stem} {'='*10}")
    
    load_status, save_status = get_status_manager(out_dir)
    status = load_status()

    # 步骤1：转录
    if status.get("step1") != "SUCCESS":
        status["step1"] = "RUNNING"
        save_status(status)
        try:
            transcript = step1_transcribe(video_path, out_dir)
            status["step1"] = "SUCCESS"
            save_status(status)
        except Exception as e:
            status["step1"] = f"FAILED: {str(e)}"
            save_status(status)
            print(f"  ✗ 步骤1失败: {e}")
            return
    else:
        print("[第1步] 视频转文字: ✓ 已完成")
        transcript = (out_dir / "01_transcript.txt").read_text(encoding="utf-8")

    # 步骤2：截图
    if status.get("step2") != "SUCCESS":
        status["step2"] = "RUNNING"
        save_status(status)
        try:
            frames = step2_extract_frames(video_path, out_dir)
            status["step2"] = "SUCCESS"
            save_status(status)
        except Exception as e:
            status["step2"] = f"FAILED: {str(e)}"
            save_status(status)
            print(f"  ✗ 步骤2失败: {e}")
            return
    else:
        print("[第2步] 截取关键帧: ✓ 已完成")
        frames = sorted((out_dir / "frames").glob("*.jpg"))

    # 步骤3：AI整理
    if status.get("step3") != "SUCCESS":
        status["step3"] = "RUNNING"
        save_status(status)
        try:
            structured = step3_structure_text(transcript, out_dir, client)
            status["step3"] = "SUCCESS"
            save_status(status)
        except Exception as e:
            status["step3"] = f"FAILED: {str(e)}"
            save_status(status)
            print(f"  ✗ 步骤3失败: {e}")
            return
    else:
        print("[第3步] AI 整理文稿结构: ✓ 已完成")
        structured = json.loads((out_dir / "02_structured.json").read_text(encoding="utf-8"))

    # 步骤4：生成考题
    if status.get("step4") != "SUCCESS":
        status["step4"] = "RUNNING"
        save_status(status)
        try:
            all_questions = step4_generate_questions(structured, out_dir, video_stem, client)
            status["step4"] = "SUCCESS"
            save_status(status)
        except Exception as e:
            status["step4"] = f"FAILED: {str(e)}"
            save_status(status)
            print(f"  ✗ 步骤4失败: {e}")
            return
    else:
        print("[第4步] AI 生成考题: ✓ 已完成")
        all_questions = json.loads((out_dir / "03_questions.json").read_text(encoding="utf-8"))

    # 步骤5：生成Word
    if status.get("step5") != "SUCCESS":
        status["step5"] = "RUNNING"
        save_status(status)
        try:
            timestamp = time.strftime("%H%M%S")
            book_path = out_dir / f"{video_stem}_教材_{timestamp}.docx"
            questions_path = out_dir / f"{video_stem}_考题_{timestamp}.docx"
            
            make_book_docx(structured, frames, book_path, video_stem)
            make_questions_docx(structured, all_questions, questions_path, video_stem)
            
            status["step5"] = "SUCCESS"
            save_status(status)
            print(f"  ✓ Word 文档生成成功")
        except Exception as e:
            status["step5"] = f"FAILED: {str(e)}"
            save_status(status)
            print(f"  ✗ 步骤5失败: {e}")
            return
    else:
        print("[第5步] 生成 Word 文档: ✓ 已完成")

    print(f"{'='*10} {video_stem} 处理完成 {'='*10}")


def main():
    print("=" * 55)
    print("  视频 → Word 产品生成器 (批量底座强化版)")
    print("=" * 55)

    video_dir = Path(VIDEO_DIR)
    video_dir.mkdir(parents=True, exist_ok=True)
    
    out_root = Path(OUTPUT_ROOT)
    out_root.mkdir(parents=True, exist_ok=True)

    # 扫描视频文件
    video_extensions = ('.mp4', '.mkv', '.avi', '.mov', '.ts')
    videos = [f for f in video_dir.iterdir() if f.suffix.lower() in video_extensions]
    
    if not videos:
        print(f"! 在目录 {VIDEO_DIR} 中未找到有效视频文件。")
        print("  请将视频放入该目录后重新运行。")
        return

    print(f"共发现 {len(videos)} 个视频待处理。")
    print(f"输出目录：{out_root}\n")

    # 初始化阿里百炼客户端
    client = anthropic.Anthropic(
        api_key=API_KEY,
        base_url=BASE_URL
    )

    for i, video_path in enumerate(videos, 1):
        # --- 中断检查点 ---
        if STOP_REQUESTED:
            print("\n🛑 检测到停止请求，正在退出...")
            break

        print(f"\n进度: [{i}/{len(videos)}]")
        process_single_video(video_path, out_root, client)

    print("\n" + "=" * 55)
    print("🎉 任务队列全部处理完成！")
    print("=" * 55)


if __name__ == "__main__":
    main()
