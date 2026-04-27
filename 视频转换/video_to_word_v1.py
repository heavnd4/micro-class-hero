"""
视频 → Word 产品生成器（单视频版）
输出：教材.docx + 考题.docx
"""

import json
import subprocess
from pathlib import Path
import anthropic
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================
# ★ 配置区（已填好，无需修改）
# ============================================================
VIDEO_PATH    = r"E:\.cc项目\视频转换\燃气改造 电回路 烟机排烟 排水路由... 【厨房水电】一口气打包全解读.mp4"
OUTPUT_DIR    = r"E:\.cc项目\视频转换\输出"
API_KEY       = "sk-sp-a057b4bf1def40ad8b44d6a908b59a84"
BASE_URL      = "https://coding.dashscope.aliyuncs.com/apps/anthropic"
MODEL         = "qwen3.5-plus"
COURSE_NAME   = "厨房水电"
FRAME_EVERY   = 60
WHISPER_MODEL = "small"
# ============================================================


def step1_transcribe(video_path: Path, out_dir: Path) -> str:
    """第一步：视频转文字"""
    txt_file = out_dir / "01_transcript.txt"

    if txt_file.exists():
        print("  ✓ 转录文稿已存在，跳过")
        return txt_file.read_text(encoding="utf-8")

    print("  → 开始转录，请耐心等待（30分钟视频约需15-25分钟）...")
    print("     转录期间可以做其他事，完成后会提示")

    from faster_whisper import WhisperModel
    model = WhisperModel(r"E:\.cc项目\models\models--Systran--faster-whisper-small\snapshots\main", device="cpu", compute_type="int8")
    segments, _ = model.transcribe(str(video_path), language="zh")

    lines = [seg.text.strip() for seg in segments]
    text = "\n".join(lines)
    txt_file.write_text(text, encoding="utf-8")
    print(f"  ✓ 转录完成，共 {len(text)} 字")
    return text


def step2_extract_frames(video_path: Path, out_dir: Path) -> list:
    """第二步：FFmpeg 截图"""
    frames_dir = out_dir / "frames"
    frames_dir.mkdir(exist_ok=True)

    existing = sorted(frames_dir.glob("*.jpg"))
    if existing:
        print(f"  ✓ 截图已存在（{len(existing)} 张），跳过")
        return existing

    print(f"  → 截图中（每 {FRAME_EVERY} 秒截一帧）...")
    output_pattern = str(frames_dir / "frame_%03d.jpg")
    cmd = [
        "ffmpeg", "-i", str(video_path),
        "-vf", f"fps=1/{FRAME_EVERY}",
        "-q:v", "2",
        output_pattern, "-y"
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, errors='ignore')
    if result.returncode != 0:
        print(f"  ✗ FFmpeg 出错：{result.stderr[-300:]}")
        return []

    frames = sorted(frames_dir.glob("*.jpg"))
    print(f"  ✓ 截图完成，共 {len(frames)} 张")
    return frames


def step3_structure_text(transcript: str, client) -> dict:
    """第三步：AI 整理文稿结构"""
    cache_file = Path(OUTPUT_DIR) / "02_structured.json"

    if cache_file.exists():
        print("  ✓ 结构化数据已存在，跳过")
        return json.loads(cache_file.read_text(encoding="utf-8"))

    print("  → 调用 AI 整理文稿，约需1-2分钟...")

    text_input = transcript[:8000] if len(transcript) > 8000 else transcript

    prompt = f"""你是一位专业的课程内容整理专家。
以下是一段技术视频的原始字幕，请完成以下任务：

1. 将内容划分为5-8个章节，每章节有简洁的标题
2. 每章节内容改写为书面化正文（去除"嗯""那个"等口语，保留所有技术术语和操作步骤）
3. 每章节正文200-400字

请用以下JSON格式返回（只返回JSON，不要其他内容）：
{{
  "title": "根据内容总结的课程标题",
  "chapters": [
    {{
      "index": 1,
      "title": "章节标题",
      "content": "书面化正文内容..."
    }}
  ]
}}

原始字幕：
{text_input}"""

    response = client.messages.create(
        model=MODEL,
        max_tokens=4000,
        messages=[{"role": "user", "content": prompt}]
    )

    # 提取文本内容，兼容多种 Block 类型（如 TextBlock, ThinkingBlock）
    raw = ""
    for block in response.content:
        if hasattr(block, "text"):
            raw += block.text
    raw = raw.strip()
    if "```" in raw:
        parts = raw.split("```")
        for part in parts:
            part = part.strip()
            if part.startswith("json"):
                part = part[4:].strip()
            if part.startswith("{"):
                raw = part
                break

    data = json.loads(raw)
    cache_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"  ✓ 整理完成，共 {len(data['chapters'])} 章")
    return data


def step4_generate_questions(structured: dict, client) -> list:
    """第四步：AI 生成考题 (一次性生成全库以保证数量和质量)"""
    cache_file = Path(OUTPUT_DIR) / "03_questions.json"

    if cache_file.exists():
        print("  ✓ 考题数据已存在，跳过")
        return json.loads(cache_file.read_text(encoding="utf-8"))

    print("  → 正在根据全文生成考题库，约需1-2分钟...")
    
    # 汇总全文内容
    full_content = "\n\n".join([f"章节{ch['index']}: {ch['title']}\n{ch['content']}" for ch in structured["chapters"]])
    
    prompt = f"""你是一位专业的教育测评专家。请根据以下课程内容，编写一套配套测试题。

课程标题：{structured.get('title', COURSE_NAME)}

要求：
1. 总题量：25-27道。
2. 题型分布：
   - 单选题：15道
   - 多选题：5道
   - 判断题：5道
   - 简答题：1-2道
3. 内容覆盖：题目必须覆盖课程的所有主要技术点（电路、插座、给排水、燃气等）。
4. 难度适中，解析详尽。

只返回一个JSON数组，格式如下：
[
  {{"type": "单选题", "question": "...", "options": ["A.", "B.", "C.", "D."], "answer": "A", "explanation": "..."}},
  {{"type": "多选题", "question": "...", "options": ["A.", "B.", "C.", "D."], "answer": "AB", "explanation": "..."}},
  {{"type": "判断题", "question": "...", "answer": "正确", "explanation": "..."}},
  {{"type": "简答题", "question": "...", "answer": "..."}}
]

课程内容：
{full_content}"""

    response = client.messages.create(
        model=MODEL,
        max_tokens=8000,
        messages=[{"role": "user", "content": prompt}]
    )

    # 提取文本内容
    raw = ""
    for block in response.content:
        if hasattr(block, "text"):
            raw += block.text
    raw = raw.strip()
    
    # 提取 JSON 数组部分
    import re
    match = re.search(r'\[.*\]', raw, re.DOTALL)
    if match:
        json_str = match.group()
        try:
            all_questions = json.loads(json_str)
        except json.JSONDecodeError:
            print("  ✗ JSON 解析失败，尝试修复并重新解析...")
            # 简单的修复逻辑：如果是因为截断导致的
            if not json_str.endswith("]"):
                json_str += "]"
            try:
                all_questions = json.loads(json_str)
            except:
                print("  ✗ 无法自动修复 JSON，请检查 API 输出")
                all_questions = []
    else:
        print("  ✗ 未在 AI 响应中找到 JSON 数组")
        all_questions = []

    cache_file.write_text(json.dumps(all_questions, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"  ✓ 考题生成完成，共 {len(all_questions)} 道题")
    return all_questions


def make_book_docx(structured: dict, frames: list, out_path: Path):
    """生成图文教材 Word"""
    doc = Document()

    section = doc.sections[0]
    section.page_width    = Inches(8.27)
    section.page_height   = Inches(11.69)
    section.left_margin   = Inches(1.18)
    section.right_margin  = Inches(1.18)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # 封面
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    run = p.add_run(structured.get("title", COURSE_NAME))
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2C, 0x4A, 0x8C)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"\n课程：{COURSE_NAME}")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()

    # 目录
    p_toc = doc.add_paragraph()
    run_toc = p_toc.add_run("目  录")
    run_toc.font.size = Pt(18)
    run_toc.font.bold = True
    doc.add_paragraph()
    for ch in structured["chapters"]:
        p = doc.add_paragraph(f"第{ch['index']}章  {ch['title']}")
        p.paragraph_format.left_indent = Pt(20)
    doc.add_page_break()

    # 正文
    n_chapters = len(structured["chapters"])
    n_frames   = len(frames)

    for ch in structured["chapters"]:
        p_title = doc.add_paragraph()
        run_title = p_title.add_run(f"第{ch['index']}章  {ch['title']}")
        run_title.font.size = Pt(16)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(0x2C, 0x4A, 0x8C)
        p_title.paragraph_format.space_before = Pt(12)
        p_title.paragraph_format.space_after  = Pt(6)

        p_body = doc.add_paragraph(ch["content"])
        p_body.paragraph_format.first_line_indent = Pt(28)

        if n_frames > 0:
            start_idx = int((ch["index"] - 1) / n_chapters * n_frames)
            end_idx   = int(ch["index"] / n_chapters * n_frames)
            chapter_frames = frames[start_idx:end_idx]

            for i, img_path in enumerate(chapter_frames[:2]):
                try:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_img = p_img.add_run()
                    run_img.add_picture(str(img_path), width=Inches(5.2))

                    p_cap = doc.add_paragraph(f"图 {ch['index']}-{i+1}  {ch['title']}")
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.runs[0].font.size = Pt(9)
                    p_cap.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                except Exception as e:
                    print(f"     插图失败（{img_path.name}）：{e}")

        doc.add_paragraph()

    doc.save(str(out_path))
    print(f"  ✓ 教材已保存 → {out_path.name}")


def make_questions_docx(structured: dict, all_questions: list, out_path: Path):
    """生成考题 Word"""
    doc = Document()

    section = doc.sections[0]
    section.page_width    = Inches(8.27)
    section.page_height   = Inches(11.69)
    section.left_margin   = Inches(1.18)
    section.right_margin  = Inches(1.18)

    # 封面
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run(f"《{structured.get('title', COURSE_NAME)}》\n配套测试题库")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2C, 0x4A, 0x8C)
    doc.add_page_break()

    q_num = 1
    all_answers = []

    # 按题型分类列出所有题目
    for qtype in ["单选题", "多选题", "判断题", "简答题"]:
        qs = [q for q in all_questions if q.get("type") == qtype]
        if not qs:
            continue

        p_type = doc.add_paragraph(f"【{qtype}】")
        p_type.runs[0].font.size = Pt(14)
        p_type.runs[0].font.bold = True
        p_type.runs[0].font.color.rgb = RGBColor(0x1A, 0x6B, 0x4A)
        doc.add_paragraph()

        for q in qs:
            p_q = doc.add_paragraph(f"{q_num}. {q['question']}")
            p_q.runs[0].font.size = Pt(11)

            if "options" in q:
                for opt in q["options"]:
                    p_opt = doc.add_paragraph(opt)
                    p_opt.paragraph_format.left_indent = Pt(28)

            all_answers.append({
                "num": q_num,
                "answer": q.get("answer", ""),
                "explanation": q.get("explanation", "")
            })
            q_num += 1
        doc.add_paragraph()

    # 答案页
    doc.add_page_break()
    p_ans_head = doc.add_paragraph("▶ 参考答案及解析")
    p_ans_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ans_head.runs[0].font.size = Pt(16)
    p_ans_head.runs[0].font.bold = True
    p_ans_head.runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    doc.add_paragraph()

    for ans in all_answers:
        ans_text = f"{ans['num']}. {ans['answer']}"
        if ans["explanation"]:
            ans_text += f"　【解析】{ans['explanation']}"
        p_a = doc.add_paragraph(ans_text)
        p_a.paragraph_format.left_indent = Pt(14)
        p_a.runs[0].font.size = Pt(10)
        p_a.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.save(str(out_path))
    print(f"  ✓ 考题已保存 → {out_path.name}")


def main():
    print("=" * 55)
    print("  视频 → Word 产品生成器")
    print("=" * 55)

    video = Path(VIDEO_PATH)
    if not video.exists():
        print(f"✗ 找不到视频文件：{VIDEO_PATH}")
        print("  请检查视频文件是否存在")
        return

    out_dir = Path(OUTPUT_DIR)
    out_dir.mkdir(parents=True, exist_ok=True)
    print(f"输出目录：{out_dir}")
    print(f"处理视频：{video.name}\n")

    # 初始化阿里百炼客户端
    client = anthropic.Anthropic(
        api_key=API_KEY,
        base_url=BASE_URL
    )

    print("[第1步] 视频转文字")
    transcript = step1_transcribe(video, out_dir)

    print("\n[第2步] 截取关键帧")
    frames = step2_extract_frames(video, out_dir)

    print("\n[第3步] AI 整理文稿结构")
    structured = step3_structure_text(transcript, client)

    print("\n[第4步] AI 生成考题")
    all_questions = step4_generate_questions(structured, client)

    print("\n[第5步] 生成 Word 文档")
    import time
    timestamp = time.strftime("%H%M%S")
    book_path      = out_dir / f"{video.stem}_教材_{timestamp}.docx"
    questions_path = out_dir / f"{video.stem}_考题_{timestamp}.docx"
    make_book_docx(structured, frames, book_path)
    make_questions_docx(structured, all_questions, questions_path)

    print("\n" + "=" * 55)
    print("🎉 全部完成！")
    print(f"   教材：{book_path}")
    print(f"   考题：{questions_path}")
    print("=" * 55)


if __name__ == "__main__":
    main()
